"""Document service for parsing and storing uploaded documents"""

import PyPDF2
from pathlib import Path
import uuid
from typing import Tuple
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from src.models.database import Document


class DocumentService:
    """Service for handling document upload and parsing"""

    def __init__(self, storage_path: str = "./uploads"):
        self.storage_path = Path(storage_path)
        self.storage_path.mkdir(exist_ok=True)

    def parse_pdf(self, file_path: str) -> Tuple[str, int]:
        """
        Extract text from PDF

        Args:
            file_path: Path to the PDF file

        Returns:
            Tuple of (extracted_text, page_count)
        """
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                page_count = len(reader.pages)

                text = ""
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"

                if not text.strip():
                    raise ValueError("No text content extracted from PDF")

                return text, page_count

        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {str(e)}")

    def parse_html(self, file_path: str) -> Tuple[str, int]:
        """
        Extract text from HTML file

        Args:
            file_path: Path to the HTML file

        Returns:
            Tuple of (extracted_text, page_count)
            Note: HTML files don't have pages, so page_count is always 1
        """
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                html_content = file.read()

            # Parse HTML and extract text
            soup = BeautifulSoup(html_content, 'lxml')

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            # Get text content
            text = soup.get_text(separator='\n\n')

            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            text = '\n'.join(line for line in lines if line)

            if not text.strip():
                raise ValueError("No text content extracted from HTML")

            return text, 1  # HTML files treated as 1 "page"

        except Exception as e:
            raise ValueError(f"Failed to parse HTML: {str(e)}")

    def parse_docx(self, file_path: str) -> Tuple[str, int]:
        """
        Extract text from DOCX file

        Args:
            file_path: Path to the DOCX file

        Returns:
            Tuple of (extracted_text, page_count)
            Note: Page count is estimated based on text length
        """
        try:
            doc = DocxDocument(file_path)

            # Extract text from all paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            text = '\n\n'.join(text_parts)

            if not text.strip():
                raise ValueError("No text content extracted from DOCX")

            # Estimate page count (rough estimate: 500 words per page, ~5 chars per word)
            estimated_pages = max(1, len(text) // 2500)

            return text, estimated_pages

        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {str(e)}")

    def save_document(self, file_data: bytes, filename: str, project_id: uuid.UUID) -> Document:
        """
        Save uploaded document to storage and parse it

        Args:
            file_data: Binary file data
            filename: Original filename
            project_id: Project UUID

        Returns:
            Document model instance
        """
        # Generate unique filename
        file_id = uuid.uuid4()
        file_extension = Path(filename).suffix
        storage_filename = f"{project_id}_{file_id}{file_extension}"
        file_path = self.storage_path / storage_filename

        # Save file to disk
        with open(file_path, 'wb') as f:
            f.write(file_data)

        # Parse text content
        text_content = ""
        page_count = 0
        parse_status = "PENDING"

        try:
            if file_extension.lower() == '.pdf':
                text_content, page_count = self.parse_pdf(str(file_path))
                parse_status = "COMPLETED"
            elif file_extension.lower() in ['.html', '.htm']:
                text_content, page_count = self.parse_html(str(file_path))
                parse_status = "COMPLETED"
            elif file_extension.lower() == '.docx':
                text_content, page_count = self.parse_docx(str(file_path))
                parse_status = "COMPLETED"
            elif file_extension.lower() == '.doc':
                parse_status = "UNSUPPORTED_FORMAT"
                text_content = "Legacy .doc format not supported. Please convert to .docx, .pdf, or .html"
                page_count = 0
            else:
                parse_status = "UNSUPPORTED_FORMAT"
                text_content = f"File format '{file_extension}' is not supported. Supported formats: PDF, HTML, DOCX"
                page_count = 0
        except Exception as e:
            parse_status = "FAILED"
            text_content = f"Parse error: {str(e)}"
            page_count = 0

        # Create document record
        document = Document(
            project_id=project_id,
            filename=filename,
            file_format=file_extension.lstrip('.').upper(),
            file_path=str(file_path),
            parse_status=parse_status,
            text_content=text_content,
            page_count=page_count
        )

        return document
